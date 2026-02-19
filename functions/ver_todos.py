import os
import streamlit as st
import pandas as pd
import io
from datetime import datetime, timedelta
from queries import buscar_expedientes
import plotly.express as px
from docx import Document
import subprocess
from pathlib import Path

BASE = Path(__file__).parent

def render(conn, catalogos):
    st.header("📊 Reportes")
    st.markdown("---")
    tab1, tab2, tab3 = st.tabs(["📊 General", "📈 Estatus"," 📅 Semanal"])
    with tab1: # GENERAL
        with st.spinner("Cargando datos..."):
            df_todos = buscar_expedientes(conn)
        
        if not df_todos.empty:
            # Filtros
            col_filt1, col_filt2, col_filt3 = st.columns(3)
            with col_filt1:
                filtro_exp = st.text_input("Filtrar por expediente:", key="filtro_exp")
            with col_filt2:
                filtro_mun = st.selectbox(
                    "Filtrar por municipio:",
                    options=["Todos"] + catalogos['municipios'],
                    key="filtro_mun"
                )
            with col_filt3:
                filtro_dep = st.selectbox(
                    "Filtrar por dependencia:",
                    options=["Todos"] + catalogos['dependencias'],
                    key="filtro_dep"
                )
            
            # Aplicar filtros
            df_filtrado = df_todos.copy()
            if filtro_exp:
                df_filtrado = df_filtrado[df_filtrado['Expediente'].str.contains(filtro_exp, na=False, case=False)]
            if filtro_mun != "Todos":
                df_filtrado = df_filtrado[df_filtrado['Municipio'] == filtro_mun]
            if filtro_dep != "Todos":
                df_filtrado = df_filtrado[df_filtrado['Dependencia'] == filtro_dep]
            
            # Mostrar
            st.dataframe(
                df_filtrado,
                use_container_width=True,
                hide_index=True,
                column_config={
                    'Expediente': st.column_config.TextColumn("Expediente", width="medium"),
                    'FechaInicio': st.column_config.TextColumn("Fecha", width="small"),
                    'Municipio': st.column_config.TextColumn("Municipio", width="medium"),
                    'Hecho': st.column_config.TextColumn("Hecho", width="large"),
                    'Dependencia': st.column_config.TextColumn("Dependencia", width="medium"),
                    'Observaciones': st.column_config.TextColumn("Observaciones", width="large")
                }
            )
            
            # Estadísticas
            st.divider()
            col_stats1, col_stats2, col_stats3 = st.columns(3)
            with col_stats1:
                st.metric("Total registros", len(df_todos))
            with col_stats2:
                st.metric("Filtrados", len(df_filtrado))
            with col_stats3:
                if len(df_filtrado) > 0:
                    hecho_comun = df_filtrado['Hecho'].mode()
                    if not hecho_comun.empty:
                        st.metric("Hecho más común", hecho_comun.iloc[0])
            
            # Botón para exportar
            if st.button("📥 Exportar a Excel", use_container_width=True, key="btn_exportar"):
                import io
                # Crear buffer en memoria
                output = io.BytesIO()
                # Escribir Excel al buffer
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df_filtrado.to_excel(writer, index=False, sheet_name='Quejas')
        
                excel_bytes = output.getvalue()
                # Botón de descarga
                st.download_button(
                    label="📥 Descargar Excel",
                    data=excel_bytes,
                    file_name=f"quejas_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="btn_descargar_excel"
                )
        else:
            st.info("No hay registros disponibles")

    with tab2: # ESTATUS
        # Filtros para el reporte
        col_filt1, col_filt2 = st.columns(2)
        
        with col_filt1:
            fecha_desde = st.text_input("Fecha desde", placeholder="DD/MM/YYYY")
        
        with col_filt2:
            fecha_hasta = st.text_input("Fecha hasta", placeholder="DD/MM/YYYY", 
                                       value=datetime.now().strftime("%d/%m/%Y"))
        
        # Generar reporte
        if st.button("📊 Generar Reporte", type="primary"):
            try:
                cursor = conn.cursor() # este funciona junto con cursor.execute()
                query = """
                SELECT 
                    Conclusión,
                    COUNT(*) as Cantidad
                FROM Expediente
                WHERE 1=1
                """
                
                # Aplicar filtros de fecha
                if fecha_desde and fecha_desde.strip():
                    fecha_desde_dt = datetime.strptime(fecha_desde, "%d/%m/%Y")
                    fecha_str = fecha_desde_dt.strftime("%d/%m/%Y")
                    query += f" AND F_Conclusion >= #{fecha_str}#"
                
                if fecha_hasta and fecha_hasta.strip():
                    fecha_hasta_dt = datetime.strptime(fecha_hasta, "%d/%m/%Y")
                    fecha_str = fecha_hasta_dt.strftime("%d/%m/%Y")
                    query += f" AND F_Conclusion <= #{fecha_str}#"
                
                query += " GROUP BY Conclusión ORDER BY COUNT(*) DESC"

                #df_reporte = pd.read_sql_query(query, conn, params=params)
                # Ejecutar query con cursor
                cursor.execute(query)

                columns = [column[0] for column in cursor.description]
                results = cursor.fetchall()
                df_reporte = pd.DataFrame.from_records(results, columns=columns)
                
                if not df_reporte.empty:
                    st.write("### Resumen por Estatus")
                    
                    # Mostrar métricas
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("Total Expedientes", df_reporte['Cantidad'].sum())
                    with col2:
                        estatus_comun = df_reporte.iloc[0]['Conclusión'] if len(df_reporte) > 0 else "N/A"
                        st.metric("Estatus Más Común", estatus_comun)
                    
                    # Mostrar tabla
                    st.dataframe(df_reporte, use_container_width=True)
                    
                    # Gráfico
                    fig = px.pie(df_reporte, values='Cantidad', names='Conclusión')
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Botón para exportar
                    csv = df_reporte.to_csv(index=False).encode('utf-8')
                    st.download_button(
                        label="📥 Descargar Reporte CSV",
                        data=csv,
                        file_name=f"reporte_estatus_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                        mime="text/csv"
                    )
                else:
                    st.info("No hay datos para el período seleccionado")
                    
            except Exception as e:
                st.error(f"Error al generar reporte: {str(e)}")

    with tab3: # SEMANAL
        if st.button("Buscar Reporte", use_container_width=True, key="btn_descargar", type="primary"):
            hoy = datetime.now()
            # 4 = viernes (0=lunes, 1=martes, ..., 6=domingo)
            dias_desde_viernes = (hoy.weekday() - 4) % 7
            if dias_desde_viernes == 0 and hoy.hour < 12:  # Si es viernes antes del mediodía
                dias_desde_viernes = 7  # Usar el viernes anterior
            ultimo_viernes = hoy - timedelta(days=dias_desde_viernes)

            try:
                fecha_str = ultimo_viernes.strftime('%d-%m-%y')
                archivo_word = f"semanal_{fecha_str}.docx"
                ruta_archivo = rf"D:\SecretariaGeneral\Informatica\ESTADISTICAS\Tarjetas\2026\{archivo_word}"
                doc = Document(ruta_archivo)
                st.info(f"📅 Ultimo viernes encontrado: {ultimo_viernes.strftime('%d/%m/%Y')}")

                st.markdown("### 📄 Vista Previa del Reporte Semanal")
                # Contenedor con borde
                with st.container():
                    st.markdown("""
                    <style>
                    .preview-box {
                        border: 2px solid #e0e0e0;
                        border-radius: 10px;
                        padding: 20px;
                        background-color: #f9f9f9;
                        max-height: 500px;
                        overflow-y: auto;
                        color: #000000;
                        text-align: center;
                    }
                    </style>
                    """, unsafe_allow_html=True)
                    preview_html = '<div class="preview-box">'
                    for i, para in enumerate(doc.paragraphs):
                        texto = para.text.strip()
                        if texto:
                            # Detectar títulos
                            if para.style.name.startswith('Heading'):
                                preview_html += f'<h3>{texto}</h3>'
                            else:
                                preview_html += f'<p>{texto}</p>'
                        
                        if i >= 10:  # Limitar vista previa
                            preview_html += '<p><i>... Ver documento completo descargándolo</i></p>'
                            break
                    preview_html += '</div>'
                    st.markdown(preview_html, unsafe_allow_html=True)
                st.markdown("---")

                # Leer el archivo
                with open(ruta_archivo, "rb") as file:
                    word_bytes = file.read()
                
                # Botón de descarga
                st.download_button(
                    label="📥 Descargar Word",
                    data=word_bytes,
                    file_name=archivo_word,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="btn_descargar_word_archivo"
                )
            except FileNotFoundError:
                st.error(f"❌ No se encontró el archivo Word especificado {archivo_word}")
            except Exception as e:
                st.error(f"❌ Error al leer el archivo: {str(e)}")

        if st.button("Generar Reporte", use_container_width=True, key="btn_generar", type="primary"):
                try:
                    ruta_script = BASE / 'reporte' / 'generar_word.py'
                    with st.spinner("Construyendo Reporte..."):
                        subprocess.run(['python', str(ruta_script)], check=True)
                    hoy = datetime.now().strftime('%d-%m-%y')
                    archivo_word = f"semanal_{hoy}.docx"
                    ruta_archivo = BASE / 'reporte' / archivo_word
                    doc = Document(ruta_archivo)
                    st.info(f"Archivo generado: {hoy}")

                    st.markdown("### 📄 Vista Previa del Reporte Semanal")
                    # Contenedor con borde
                    with st.container():
                        st.markdown("""
                        <style>
                        .preview-box {
                            border: 2px solid #e0e0e0;
                            border-radius: 10px;
                            padding: 20px;
                            background-color: #f9f9f9;
                            max-height: 500px;
                            overflow-y: auto;
                            color: #000000;
                            text-align: center;
                        }
                        </style>
                        """, unsafe_allow_html=True)
                        preview_html = '<div class="preview-box">'
                        for i, para in enumerate(doc.paragraphs):
                            texto = para.text.strip()
                            if texto:
                                # Detectar títulos
                                if para.style.name.startswith('Heading'):
                                    preview_html += f'<h3>{texto}</h3>'
                                else:
                                    preview_html += f'<p>{texto}</p>'
                            
                            if i >= 10:  # Limitar vista previa
                                preview_html += '<p><i>... Ver documento completo descargándolo</i></p>'
                                break
                        preview_html += '</div>'
                        st.markdown(preview_html, unsafe_allow_html=True)
                    st.markdown("---")

                    # Leer el archivo
                    with open(ruta_archivo, "rb") as file:
                        word_bytes = file.read()
                    
                    # Botón de descarga
                    st.download_button(
                        label="📥 Descargar Word",
                        data=word_bytes,
                        file_name=archivo_word,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="btn_descargar_word_generado"
                    )
                except FileNotFoundError:
                    st.error(f"❌ No se encontró el archivo Word especificado {archivo_word}")
                except Exception as e:
                    st.error(f"❌ Error al leer el archivo: {str(e)}")