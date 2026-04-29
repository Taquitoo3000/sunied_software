import os
import streamlit as st
import pandas as pd
import io
from datetime import datetime, timedelta
from queries import buscar_expedientes, busqueda_personalizada
import plotly.express as px
from docx import Document
import subprocess
from pathlib import Path
from functions.reporte import predecir_rec
from functions.reporte.generar_reporte_pred import oraculo_reporte
from sqlalchemy import text

BASE = Path(__file__).parent

def render(conn, catalogos):
    st.header("📊 Reportes")
    st.markdown("---")
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 General", "📈 Estatus","📅 Semanal", "⭐ Personalizado", "🔮 Prediccion"])
    with tab1: # GENERAL
        with st.spinner("Cargando datos..."):
            df_todos = buscar_expedientes(conn)
        
        if not df_todos.empty:
            # Filtros
            col_filt1, col_filt2, col_filt3, col_filt4, col_filt5 = st.columns(5)
            with col_filt1:
                filtro_est = st.selectbox(
                    "Filtrar por Estatus:",
                    options=["Todos"] + catalogos['Status'],
                    key="filtro_est"
                )
            with col_filt2:
                fecha_desde = st.date_input("Fecha desde", key="fecha_desde_1")
            with col_filt3:
                fecha_hasta = st.date_input("Fecha hasta", key="fecha_hasta_1")
            with col_filt4:
                filtro_mun = st.selectbox(
                    "Filtrar por municipio:",
                    options=["Todos"] + catalogos['municipios'],
                    key="filtro_mun"
                )
            with col_filt5:
                filtro_dep = st.selectbox(
                    "Filtrar por dependencia:",
                    options=["Todos"] + catalogos['dependencias'],
                    key="filtro_dep"
                )
            
            # Aplicar filtros
            df_filtrado = df_todos.copy()
            df_filtrado['FechaInicio'] = pd.to_datetime(df_filtrado['FechaInicio'],format="%d/%m/%Y", errors='coerce')
            df_filtrado['F_Conclusion'] = pd.to_datetime(df_filtrado['F_Conclusion'],format="%d/%m/%Y", errors='coerce')
            if fecha_desde and fecha_hasta:
                df_filtrado = df_filtrado[
                    ((df_filtrado['FechaInicio'] >= pd.to_datetime(fecha_desde,format="%d/%m/%Y",errors="coerce")) &
                    (df_filtrado['FechaInicio'] <= pd.to_datetime(fecha_hasta,format="%d/%m/%Y",errors="coerce")))
                ]
            if filtro_est != "Todos":
                df_filtrado = df_filtrado[df_filtrado['Conclusión']==filtro_est]
            if filtro_mun != "Todos":
                df_filtrado = df_filtrado[df_filtrado['Municipio'] == filtro_mun]
            if filtro_dep != "Todos":
                df_filtrado = df_filtrado[df_filtrado['Dependencia'] == filtro_dep]
            
            # Mostrar
            st.dataframe(
                df_filtrado,
                width='stretch',
                hide_index=True,
                column_config={
                    'Expediente': st.column_config.TextColumn("Expediente", width="small"),
                    'FechaInicio': st.column_config.DateColumn("Fecha", width="small",format='DD/MM/YYYY'),
                    'F_Conclusion': st.column_config.DateColumn("Fecha Final", width="small",format='DD/MM/YYYY'),
                    'Municipio': st.column_config.TextColumn("Municipio", width="medium"),
                    'Hecho': st.column_config.TextColumn("Hecho", width="medium"),
                    'DireccionMunicipal': st.column_config.TextColumn("Autoridad", width="medium"),
                    'Dependencia': st.column_config.TextColumn("Dependencia", width="medium")
                }
            )
            
            # Estadísticas
            st.divider()
            col_stats1, col_stats2, col_stats3 = st.columns(3)
            with col_stats1:
                st.metric("Total registros", len(df_todos))
            with col_stats2:
                st.metric("Filtrados", len(df_filtrado))
            with col_stats3:
                if len(df_filtrado) > 0:
                    hecho_comun = df_filtrado['Hecho'].mode()
                    if not hecho_comun.empty:
                        st.metric("Hecho más común", hecho_comun.iloc[0])
            excel_bytes = descargar_excel(df_filtrado)
            # Botón de descarga
            st.download_button(
                label="Descargar Excel",
                data=excel_bytes,
                file_name=f"quejas_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="btn_descargar_excel",
                icon=":material/download:",
            )
        else:
            st.info("No hay registros disponibles")

    with tab2: # ESTATUS
        # Filtros para el reporte
        col_filt1, col_filt2 = st.columns(2)
        
        with col_filt1:
            fecha_desde = st.text_input("Fecha desde", placeholder="DD/MM/YYYY")
        
        with col_filt2:
            fecha_hasta = st.text_input("Fecha hasta", placeholder="DD/MM/YYYY", 
                                       value=datetime.now().strftime("%d/%m/%Y"))
        
        # Generar reporte
        if st.button("📊 Generar Reporte", type="primary"):
            try:
                
                # Aplicar filtros de fecha
                if fecha_desde and fecha_desde.strip():
                    fecha_desde_dt = datetime.strptime(fecha_desde, "%d/%m/%Y")
                
                if fecha_hasta and fecha_hasta.strip():
                    fecha_hasta_dt = datetime.strptime(fecha_hasta, "%d/%m/%Y")

                query = """
                SELECT 
                    Conclusión,
                    COUNT(*) as Cantidad
                FROM Expediente
                WHERE 1=1
                    AND F_Conclusion >= :fecha_desde_dt
                    AND F_Conclusion <= :fecha_hasta_dt
                GROUP BY Conclusión
                ORDER BY COUNT(*) DESC
                """

                #df_reporte = pd.read_sql_query(query, conn, params=params)
                # Ejecutar query con cursor
                df_reporte = pd.read_sql_query(text(query), conn, params={
                    'fecha_desde_dt': fecha_desde_dt,
                    'fecha_hasta_dt': fecha_hasta_dt
                })
                
                if not df_reporte.empty:
                    st.write("### Resumen por Estatus")
                    
                    # Mostrar métricas
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("Total Expedientes", df_reporte['Cantidad'].sum())
                    with col2:
                        estatus_comun = df_reporte.iloc[0]['Conclusión'] if len(df_reporte) > 0 else "N/A"
                        st.metric("Estatus Más Común", estatus_comun)
                    
                    # Mostrar tabla
                    st.dataframe(df_reporte, width='stretch')
                    
                    # Gráfico
                    fig = px.pie(df_reporte, values='Cantidad', names='Conclusión')
                    st.plotly_chart(fig, width='stretch')
                else:
                    st.info("No hay datos para el período seleccionado")
            except Exception as e:
                st.error(f"Error al generar reporte: {str(e)}")

    with tab3: # SEMANAL
        semanales = os.listdir(BASE / 'reporte')
        semanal_seleccionado = st.selectbox(
            "Selecciona un reporte semanal:",
            options=semanales
        )
        if st.button("Buscar Reporte", width='stretch', key="btn_descargar", type="primary"):
            with st.spinner("Cargando reporte..."):
                try:
                    archivo_word = semanal_seleccionado
                    ruta_archivo = os.path.join(
                        r"D:/SecretariaGeneral/Informatica/ESTADISTICAS/Tarjetas/2026",
                        archivo_word
                    )
                    doc = Document(ruta_archivo)
                    st.markdown("### 📄 Vista Previa del Reporte Semanal")
                    # Contenedor con borde
                    with st.container():
                        st.markdown("""
                        <style>
                        .preview-box {
                            border: 2px solid #e0e0e0;
                            border-radius: 10px;
                            padding: 20px;
                            background-color: #f9f9f9;
                            max-height: 500px;
                            overflow-y: auto;
                            color: #000000;
                            text-align: center;
                        }
                        </style>
                        """, unsafe_allow_html=True)
                        preview_html = '<div class="preview-box">'
                        for i, para in enumerate(doc.paragraphs):
                            texto = para.text.strip()
                            if texto:
                                # Detectar títulos
                                if para.style.name.startswith('Heading'):
                                    preview_html += f'<h3>{texto}</h3>'
                                else:
                                    preview_html += f'<p>{texto}</p>'
                            
                            if i >= 10:  # Limitar vista previa
                                preview_html += '<p><i>... Ver documento completo descargándolo</i></p>'
                                break
                        preview_html += '</div>'
                        st.markdown(preview_html, unsafe_allow_html=True)
                    st.markdown("---")

                    # Leer el archivo
                    with open(ruta_archivo, "rb") as file:
                        word_bytes = file.read()
                    
                    # Botón de descarga
                    st.download_button(
                        label="📥 Descargar Word",
                        data=word_bytes,
                        file_name=archivo_word,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        width='stretch',
                        key="btn_descargar_word_archivo"
                    )
                except FileNotFoundError:
                    st.error(f"❌ No se encontró el archivo Word especificado {archivo_word}")
                except Exception as e:
                    st.error(f"❌ Error al leer el archivo: {str(e)}")

        if st.button("Generar Reporte", width='stretch', key="btn_generar", type="primary"):
                try:
                    ruta_script = BASE / 'reporte' / 'generar_word.py'
                    with st.spinner("Construyendo Reporte..."):
                        subprocess.run(['python', str(ruta_script)], check=True)
                    hoy = datetime.now().strftime('%d-%m-%y')
                    archivo_word = f"semanal_{hoy}.docx"
                    ruta_archivo = BASE / 'reporte' / archivo_word
                    doc = Document(ruta_archivo)
                    st.info(f"Archivo generado: {hoy}")

                    st.markdown("### 📄 Vista Previa del Reporte Semanal")
                    # Contenedor con borde
                    with st.container():
                        st.markdown("""
                        <style>
                        .preview-box {
                            border: 2px solid #e0e0e0;
                            border-radius: 10px;
                            padding: 20px;
                            background-color: #f9f9f9;
                            max-height: 500px;
                            overflow-y: auto;
                            color: #000000;
                            text-align: center;
                        }
                        </style>
                        """, unsafe_allow_html=True)
                        preview_html = '<div class="preview-box">'
                        for i, para in enumerate(doc.paragraphs):
                            texto = para.text.strip()
                            if texto:
                                # Detectar títulos
                                if para.style.name.startswith('Heading'):
                                    preview_html += f'<h3>{texto}</h3>'
                                else:
                                    preview_html += f'<p>{texto}</p>'
                            
                            if i >= 10:  # Limitar vista previa
                                preview_html += '<p><i>... Ver documento completo descargándolo</i></p>'
                                break
                        preview_html += '</div>'
                        st.markdown(preview_html, unsafe_allow_html=True)
                    st.markdown("---")

                    # Leer el archivo
                    with open(ruta_archivo, "rb") as file:
                        word_bytes = file.read()
                    
                    # Botón de descarga
                    st.download_button(
                        label="📥 Descargar Word",
                        data=word_bytes,
                        file_name=archivo_word,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        width='stretch',
                        key="btn_descargar_word_generado"
                    )
                except FileNotFoundError:
                    st.error(f"❌ No se encontró el archivo Word especificado {archivo_word}")
                except Exception as e:
                    st.error(f"❌ Error al leer el archivo: {str(e)}")

    with tab4: # PERSONALIZADO
        carpeta = r"C:\Users\PRODHEG\Desktop\isael\sql_querys"
        archivos = os.listdir(carpeta)
        archivo_seleccionado = st.selectbox(
            "Selecciona una query:",
            options=archivos
        )

        st.write(f"Seleccionaste: {archivo_seleccionado}")

        ruta_completa = os.path.join(carpeta, archivo_seleccionado)
        with st.expander("Ver contenido del query"):
            with open(ruta_completa, 'r', encoding='utf-8') as f:
                contenido = f.read()
                st.code(contenido, language='sql')

        if st.button("Ejecutar Query", use_container_width=True, key="btn_personalizado", type="primary"):
            try:
                with st.spinner("Cargando datos..."):
                    #conn2=get_connection()
                    #df_reporte = busqueda_personalizada(conn2,archivo_seleccionado)
                    st.session_state["df_reporte"] = df_reporte
                    #df_reporte = pd.DataFrame.from_records(df_reporte, columns=columns)
                    if "df_reporte" in st.session_state:
                        df_reporte = st.session_state["df_reporte"]

                        st.write("### Resultado")
                        st.dataframe(df_reporte, width='stretch')

                        xls_bytes = descargar_excel(df_reporte)
                        st.download_button(
                            label="📥 Descargar Excel",
                            data=xls_bytes,
                            file_name=f"quejas_{datetime.now().strftime('%d%m%Y_%H%M')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            width='stretch',
                            key="btn_descargar_xls2"
                        )
                    
            except Exception as e:
                st.error(f"Error al generar reporte: {str(e)}")

    with tab5: # PREDICCION
        col_pred1, col_pred2 = st.columns(2)
        with col_pred1:
            expediente = st.text_input("Número de Expediente", placeholder="Ej: 0002-2024-A")
            if st.button("Consultar", width='stretch'):
                if expediente:
                    with st.spinner("Consultando modelo..."):
                        try:
                            resultado=predecir_rec.predecir(expediente,conn)
                            if "error" in resultado:
                                st.error(resultado["error"])
                            else:
                                st.metric("Probabilidad", f"{resultado['probabilidad_recomendacion']}%")
                                st.metric("Predicción", resultado['prediccion'])
                                st.metric("Riesgo", resultado['riesgo'])
                        except Exception as e:
                            st.error(f"Error al conectar con el modelo")
                else:
                    st.warning("Por favor ingresa un número de expediente.")
        with col_pred2:
            st.write("### Reporte de Predicciones")
            if st.button("Generar Reporte"):
                with st.spinner("Generando reporte..."):
                    try:
                        df=oraculo_reporte(conn)
                        st.write(f"Total expedientes analizados: {len(df)}")
                        st.write(f"Riesgo Alto: {(df['riesgo'] == 'Alto').sum()} | Riesgo Medio: {(df['riesgo'] == 'Medio').sum()} | Riesgo Bajo: {(df['riesgo'] == 'Bajo').sum()}")
                        st.dataframe(df, width='stretch')
                    except Exception as e:
                        st.error(f"Error al generar el reporte: {str(e)}")

def descargar_excel(df):
    import io
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Quejas')
    excel_bytes = output.getvalue()
    return excel_bytes